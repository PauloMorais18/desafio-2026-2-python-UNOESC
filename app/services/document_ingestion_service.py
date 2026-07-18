"""Extract uploaded documents and index their text in the knowledge table."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy import func, select, update
from sqlalchemy.orm import Session

from app.models.conhecimento import Knowledge
from app.services.embedding_service import EmbeddingService, EmbeddingUnavailableError

logger = logging.getLogger(__name__)


class DocumentIngestionError(ValueError):
    """Raised when a document cannot provide searchable textual content."""


def extract_text(path: Path) -> str:
    """Extract text from one of the document formats accepted by the API."""
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8-sig")
        if suffix == ".pdf":
            return "\n\n\f\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        if suffix == ".docx":
            document = Document(path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise DocumentIngestionError(f"Não foi possível ler o arquivo: {exc}") from exc
    raise DocumentIngestionError("Formato não suportado para extração de texto.")


def split_text(text: str, chunk_size: int = 1600, overlap: int = 250) -> list[str]:
    """Split by PDF page, remove repeated boilerplate, then create bounded chunks."""
    raw_pages = text.split("\f")
    has_page_markers = len(raw_pages) > 1
    seen_sentences: set[str] = set()
    sections: list[str] = []

    for page_number, raw_page in enumerate(raw_pages, start=1):
        compact_page = re.sub(r"\s+", " ", raw_page).strip()
        if not compact_page:
            continue
        unique_sentences = []
        for sentence in re.split(r"(?<!\d)(?<=[.!?])\s+", compact_page):
            sentence = sentence.strip()
            normalized_sentence = re.sub(r"\s+", " ", sentence).casefold()
            if not normalized_sentence or normalized_sentence in seen_sentences:
                continue
            seen_sentences.add(normalized_sentence)
            unique_sentences.append(sentence)
        if unique_sentences:
            page_text = " ".join(unique_sentences)
            if has_page_markers:
                page_text = f"Página {page_number}\n{page_text}"
            sections.append(page_text)

    normalized = "\n\n".join(sections).strip()
    if not normalized:
        return []
    chunks: list[str] = []
    for section in sections:
        start = 0
        while start < len(section):
            end = min(start + chunk_size, len(section))
            if end < len(section):
                sentence = section.rfind(". ", start, end)
                if sentence > start + chunk_size // 2:
                    end = sentence + 2
            chunk = section[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(section):
                break
            start = max(start + 1, end - overlap)
    return chunks


class DocumentIngestionService:
    """Persist document chunks so the normal assistant retrieval can find them."""

    def __init__(self, session: Session) -> None:
        self.session = session

    def has_active_chunks(self, document_name: str) -> bool:
        """Return whether a document is already available to retrieval."""
        total = self.session.scalar(
            select(func.count(Knowledge.id)).where(
                Knowledge.source_document == document_name,
                Knowledge.active.is_(True),
            )
        )
        return bool(total)

    def index_missing(self, directory: Path, extensions: set[str]) -> dict[str, int]:
        """Index bundled documents once, without duplicating active chunks."""
        directory.mkdir(parents=True, exist_ok=True)
        indexed: dict[str, int] = {}
        for path in sorted(directory.iterdir()):
            if (
                path.is_file()
                and path.suffix.lower() in extensions
                and not self.has_active_chunks(path.name)
            ):
                indexed[path.name] = self.index(path)
        return indexed

    def index(self, path: Path) -> int:
        """Replace the active index for a document and return its chunk count."""
        chunks = split_text(extract_text(path))
        if not chunks:
            raise DocumentIngestionError(
                "O arquivo não contém texto pesquisável. PDFs digitalizados precisam de OCR."
            )
        try:
            embeddings: list[list[float] | None] = EmbeddingService().embed_documents(chunks)
        except EmbeddingUnavailableError as exc:
            logger.warning("Documento indexado sem embeddings: %s", exc)
            embeddings = [None] * len(chunks)
        self.session.execute(
            update(Knowledge)
            .where(Knowledge.source_document == path.name, Knowledge.active.is_(True))
            .values(active=False)
        )
        for index, (content, embedding) in enumerate(zip(chunks, embeddings, strict=True), start=1):
            self.session.add(Knowledge(
                title=f"{path.stem} — trecho {index}", content=content,
                category="documento", source_document=path.name,
                chunk_index=index, embedding=embedding, active=True,
            ))
        self.session.commit()
        return len(chunks)

    def deactivate(self, document_name: str) -> None:
        """Remove a document from future retrieval without breaking audit logs."""
        self.session.execute(
            update(Knowledge)
            .where(Knowledge.source_document == document_name, Knowledge.active.is_(True))
            .values(active=False)
        )
        self.session.commit()
